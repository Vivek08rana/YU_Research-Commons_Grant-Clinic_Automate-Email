from docx import Document
import pandas as pd
from tkinter import *
from tkinter import filedialog

"""
By: Vivek Rana
Usage: York University - Research Commons, Grant Clinic
About: This Python program reads the Grant Clinic Tracking Excel Spreadsheet to generate
unique emails for each peer reviewer by reading the professor name, grant name, and deadlines 
from the spreadsheet.
"""


# The function that reads the Excel spreadsheet and generates email templates in the form of Word documents
def main(excel_path_var, word_path_var_1, word_path_var_2, output_path_var, op_manager_entry_var, status_box):
    # Set properties of status box used to communicate errors and success to users
    status_box.config(state='normal')
    status_box.delete('1.0', END)

    # Reads the Excel spreadsheet if one is selected. Displays an error if user did not select Excel file
    try:
        xls = pd.ExcelFile(excel_path_var.get())
        df1 = pd.read_excel(xls, 'Participants (Reviewers)')
    except FileNotFoundError:
        status_box.insert(END, "ERROR: Please select a Excel File")
        return  # Return nothing, stops the rest of the function from executing, preventing errors

    # Store the info of the columns of interest
    participant_names = df1['Name']
    grant_names = df1['Grant']

    nda_check = df1['Send_Peer_Reviewer_NDA']
    evaluations_returned = df1["Evaluation_Forms_returned"].astype(str).str.strip("()")

    print(evaluations_returned)

    # Merge so column and row data can be read appropriately
    merge1 = pd.merge(
        participant_names,
        grant_names,
        left_index=True,
        right_index=True
    )

    merge2 = pd.merge(
        nda_check,
        evaluations_returned,
        left_index=True,
        right_index=True
    )
    combination_of_all = pd.merge(
        merge1,
        merge2,
        left_index=True,
        right_index=True
    )
    print(combination_of_all)

    # Check if word documents were selected. Displays an error if user did not select two Word files
    if word_path_var_1.get() == '' or word_path_var_2.get() == '':
        status_box.insert(END, "ERROR: Please select a Word File")
        return  # Return nothing, stops the rest of the function from executing, preventing errors

    for row in combination_of_all.itertuples():
        # Check if the participant wants an NDA to dictate Word template for email
        if row.Send_Peer_Reviewer_NDA == "x":
            template_file_path = word_path_var_2.get()
        else:
            template_file_path = word_path_var_1.get()

        # Reads the Excel sheet to fetch the info that we want to replace in the Word template
        # Stores the info of interest in variables
        variables = {
            "${PARTICIPANT_NAME}": row.Name,
            "${REVIEWER_NAME}": "[REVIEWER NAME HERE]",
            "${GRANT_NAME}": row.Grant,
            "${EVAL_RETURNED}": row.Evaluation_Forms_returned,
            "${OPERATIONS_MANAGER}": op_manager_entry_var.get(),
        }

        if row.Evaluation_Forms_returned == "nan":
            variables = {
                "${PARTICIPANT_NAME}": row.Name,
                "${REVIEWER_NAME}": "[REVIEWER NAME HERE]",
                "${GRANT_NAME}": row.Grant,
                "${EVAL_RETURNED}": "[RETURN DATE HERE]",
                "${OPERATIONS_MANAGER}": op_manager_entry_var.get(),
            }

        # Assigns output file names and the template Word file that should be read
        output_file_path = output_path_var.get() + '\\' + f'{row.Name} - Email.docx'
        template_document = Document(template_file_path)

        # Replacement process. Checks the variables assigned above and replaces the variables in the Word template
        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        # Saves the final result as a new generated document in the selected output path
        # If no output folder is selected, displays an error
        try:
            template_document.save(output_file_path)
            status_box.insert(END, (row, ' file generated\n'))
        except PermissionError:
            status_box.insert(END, "ERROR: Please select a output folder")
            break


# Method to replace the variables in the Word template
# using the variables containing the information of interest from the spreadsheet
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


# The function that generates the GUI
def main2():
    # Set properties of Window
    root = Tk()
    root.title('Grant Draft Application Email Template Generator')
    root.iconbitmap('YU_Icon.ico')

    # General instructions
    guide_label = Label(root, text="Welcome! This program will help you generate Email templates for sending out "
                                   "Grant Draft Applications to reviewers."
                                   "\n\nStep 1: Select the Grant Clinic Tracking Excel File"
                                   "\nStep 2: Select the Word file for the Email Template WITHOUT NDA"
                                   "\nStep 3: Select the Word file for the Email Template WITH NDA"
                                   "\nStep 4: Select the output folder"
                                   "\nStep 5: Enter the name of the Operations Manager"
                                   "\nStep 6: Press generate!")

    # Button that user can press to execute method to select Excel file
    excel_label = Label(root, text="\n\nExcel File: ")
    excel_path_var = StringVar(root, value='')
    excel_path_entry = Entry(root, textvariable=excel_path_var, state='disabled', width=60)
    excel_button = Button(
        root, text="Select", bg='white', width=20, command=lambda: excel_selection(root, excel_path_var)
    )

    # Button that user can press to execute method to select Word files
    word_label_1 = Label(root, text="\n\nWord File w/o NDA: ")
    word_path_var_1 = StringVar(root, value='')
    word_path_entry_1 = Entry(root, textvariable=word_path_var_1, state='disabled', width=60)
    word_button_1 = Button(
        root, text="Select", bg='white', width=20, command=lambda: word_selection(root, word_path_var_1)
    )

    word_label_2 = Label(root, text="\n\nWord File w/NDA: ")
    word_path_var_2 = StringVar(root, value='')
    word_path_entry_2 = Entry(root, textvariable=word_path_var_2, state='disabled', width=60)
    word_button_2 = Button(
        root, text="Select", bg='white', width=20, command=lambda: word_selection2(root, word_path_var_2)
    )

    # Button that user can press to execute method to select Output Folder
    output_label = Label(root, text="\n\nOutput Folder: ")
    output_path_var = StringVar(root, value='')
    output_path_entry = Entry(root, textvariable=output_path_var, state='disabled', width=60)
    output_button = Button(
        root, text="Select", bg='white', width=20, command=lambda: output_selection(root, output_path_var)
    )

    # Input field that allows for the name of the Operations Manager because it will change in the future
    op_manager_label = Label(root, text="\n\nOperations Manager: ")
    op_manager_entry_var = StringVar(root, value='Av')
    op_manager_entry = Entry(root, textvariable=op_manager_entry_var, width=60)

    # Button that user can press to generate the email templates using the files they have selected
    generate_button = Button(
        root, text="Generate", bg='white', width=30,
        command=lambda: main(excel_path_var, word_path_var_1, word_path_var_2,
                             output_path_var, op_manager_entry, status_box)
    )

    # Sets the properties of the status box used to communicate errors and success to users
    status_box = Text(root, height=10, width=100, bg="light cyan", state='disabled')

    # GUI Grid System
    guide_label.grid(row=0, column=1)

    excel_label.grid(row=3, column=0)
    excel_path_entry.grid(row=3, column=1)
    excel_button.grid(row=3, column=3)

    word_label_1.grid(row=5, column=0)
    word_path_entry_1.grid(row=5, column=1)
    word_button_1.grid(row=5, column=3)

    word_label_2.grid(row=7, column=0)
    word_path_entry_2.grid(row=7, column=1)
    word_button_2.grid(row=7, column=3)

    output_label.grid(row=9, column=0)
    output_path_entry.grid(row=9, column=1)
    output_button.grid(row=9, column=3)

    op_manager_label.grid(row=11, column=0)
    op_manager_entry.grid(row=11, column=1)

    generate_button.grid(row=12, column=1)

    status_box.grid(row=13, column=1)

    root.mainloop()


# Allow the user to select an Excel spreadsheet
def excel_selection(root, excel_path_var):
    excel_filename = filedialog.askopenfilename(
        parent=root, title='Select the Excel File', filetypes=[("xlsx files", "*.xlsx")]
    )
    excel_path_var.set(excel_filename)


# Allow the user to select a Word file template without NDA
def word_selection(root, word_path_var_1):
    word_filename = filedialog.askopenfilename(
        parent=root, title='Select the Word File', filetypes=[("docx files", "*.docx")]
    )
    word_path_var_1.set(word_filename)


# Allow the user to select a Word file template with NDA
def word_selection2(root, word_path_var_2):
    word_filename = filedialog.askopenfilename(
        parent=root, title='Select the Word File', filetypes=[("docx files", "*.docx")]
    )
    word_path_var_2.set(word_filename)


# Allows the user to select an output folder
def output_selection(root, output_path_var):
    output_folder_name = filedialog.askdirectory(parent=root, title='Select the Output Folder')
    output_path_var.set(output_folder_name)


# Executes the function that creates the GUI
if __name__ == '__main__':
    main2()
